"""
图片转 Word 排版工具 v2.4
- 标题文字可编辑（默认"估价对象现状照片"）
- 布局选择：2×2 / 2×3 / 2×4
- 表格边框可隐藏
- 拖拽添加图片（tkinterdnd2）
- 修复拖拽区域尺寸 + 改用 pack 布局避免 DnD 事件路由问题
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os, sys
from PIL import Image, ImageTk

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
except ImportError:
    HAS_DND = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def set_cell_width(cell, width_in):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(int(width_in * 1440)))
    tc_w.set(qn('w:type'), 'dxa')
    tc_pr.append(tc_w)


def add_cell_image(cell, img_path, width_in):
    """Convert image to clean JPEG via temp file to bypass EXIF/stream bugs in python-docx."""
    import tempfile
    from PIL import Image as PILImage
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    with PILImage.open(img_path) as pil_img:
        # Ensure RGB (handles palette, CMYK, RGBA, etc.)
        if pil_img.mode in ('P', 'LA', 'PA', 'CMYK', 'RGBA'):
            clean_img = pil_img.convert('RGBA')
            background = PILImage.new('RGBA', clean_img.size, (255, 255, 255, 255))
            background.alpha_composite(clean_img)
            clean_img = background.convert('RGB')
        elif pil_img.mode != 'RGB':
            clean_img = pil_img.convert('RGB')
        else:
            clean_img = pil_img
        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.jpg')
        os.close(tmp_fd)
        try:
            # Save as JPEG to match original template format
            clean_img.save(tmp_path, format='JPEG', quality=95)
            run.add_picture(tmp_path, width=Inches(width_in))
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)


def add_cell_text(cell, text, size_pt, rgb, align, bold=False):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*rgb)
    run.font.bold = bold


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    # Remove existing tblBorders
    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    # Add no-border tblBorders
    tbl_borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)


class ImageToWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("图片排版工具 v2.4")
        self.root.geometry("1020x720")
        self.root.minsize(900, 620)
        self.root.configure(bg="#f0f0f0")

        self.image_paths = []
        self.current_idx = -1
        self._preview_img = None

        # 设置变量
        self.caption_tpl   = tk.StringVar(value="{f}")
        self.caption_size  = tk.IntVar(value=10)
        self.caption_align = tk.StringVar(value="居中")
        self.output_name   = tk.StringVar(value="图片排版文档")
        self.doc_title     = tk.StringVar(value="估价对象现状照片")
        self.show_border   = tk.BooleanVar(value=True)

        if not HAS_DOCX:
            messagebox.showerror("缺少依赖", "请先安装：\npip install python-docx Pillow")
            sys.exit(1)

        self._build_ui()
        self.root.after(100, self._init_preview)

    # ── UI ────────────────────────────────────────────────────
    def _build_ui(self):
        top_bar = tk.Frame(self.root, bg="#2d5a8a", height=44)
        top_bar.pack(fill="x")
        top_bar.pack_propagate(False)
        tk.Label(top_bar, text="图片排版工具 v2.4", font=("微软雅黑", 14, "bold"),
                 fg="white", bg="#2d5a8a").pack(side="left", padx=16, pady=8)
        tk.Label(top_bar, text="匹配照片模板：图片+标注交替行，2列布局",
                 font=("微软雅黑", 9), fg="#aaccff", bg="#2d5a8a").pack(side="left", pady=8)

        # 用普通 pack 布局代替 ttk.PanedWindow，避免 DnD 事件路由问题
        body = tk.Frame(self.root, bg="#f0f0f0")
        body.pack(fill="both", expand=True, pady=(4, 0))

        # ── 左侧：图片列表（固定宽度 270） ──
        left = tk.Frame(body, width=270, bg="#f0f0f0")
        left.pack(side="left", fill="both", padx=(4, 2), pady=0)
        left.pack_propagate(False)

        # 拖拽区域：大一点（80px），用 pack_propagate 保持尺寸
        self.drop_frame = tk.Frame(left, bg="#e8f4fd", bd=2, relief="solid")
        self.drop_frame.pack(fill="x", padx=4, pady=(4, 4))
        self.drop_frame.pack_propagate(False)
        self.drop_frame.configure(height=80)
        self.drop_label = tk.Label(self.drop_frame,
                                   text="将图片或文件夹拖拽到此处" if HAS_DND else "点击选择图片文件",
                                   font=("微软雅黑", 10, "bold"), bg="#e8f4fd", fg="#2d5a8a",
                                   cursor="hand2")
        self.drop_label.pack(fill="both", expand=True, padx=8, pady=12)
        self._setup_drop_hint()

        card = tk.Frame(left, bg="white", bd=1, relief="solid")
        card.pack(fill="both", expand=True, padx=4, pady=(0, 4))
        tk.Label(card, text="已添加的图片", font=("微软雅黑", 10, "bold"),
                 bg="white", fg="#2d5a8a").pack(anchor="w", padx=12, pady=(10, 2))
        tk.Label(card, text="双击打开  |  Ctrl+单击多选",
                 font=("微软雅黑", 8), bg="white", fg="#999").pack(anchor="w", padx=12, pady=(0, 6))

        lf = tk.Frame(card, bg="white")
        lf.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        sy = ttk.Scrollbar(lf); sy.pack(side="right", fill="y")
        sx = ttk.Scrollbar(lf, orient="horizontal"); sx.pack(side="bottom", fill="x")
        self.listbox = tk.Listbox(lf, font=("微软雅黑", 10), bg="white",
            yscrollcommand=sy.set, xscrollcommand=sx.set,
            selectmode=tk.EXTENDED, activestyle="none", relief="flat", highlightthickness=0, bd=0)
        self.listbox.pack(side="left", fill="both", expand=True)
        sy.config(command=self.listbox.yview); sx.config(command=self.listbox.xview)
        self.listbox.bind("<<ListboxSelect>>", self._on_select)
        self.listbox.bind("<Double-Button-1>", lambda e: self._open_img())

        btn_box = tk.Frame(left, bg="#f0f0f0")
        btn_box.pack(fill="x", padx=4, pady=(0, 4))
        for txt, cmd, col in [
            ("添加图片",   self._add_images,   "#217346"),
            ("移除选中",   self._remove_sel,   "#c42b1c"),
            ("上移",      lambda: self._shift(-1), "#777"),
            ("下移",      lambda: self._shift(1),  "#777"),
            ("清空全部",   self._clear_all,    "#999"),
        ]:
            b = tk.Button(btn_box, text=txt, command=cmd, font=("微软雅黑", 9),
                          bg=col, fg="white", relief="flat", cursor="hand2")
            b.pack(side="left" if "清空" not in txt else "right",
                   padx=2, pady=4, fill="x", expand=True if "清空" not in txt else False)

        # ── 中间：预览（可伸缩） ──
        mid = tk.Frame(body, bg="#f0f0f0")
        mid.pack(side="left", fill="both", expand=True, padx=2, pady=0)
        card_m = tk.Frame(mid, bg="white", bd=1, relief="solid")
        card_m.pack(fill="both", expand=True, padx=4, pady=8)
        tk.Label(card_m, text="图片预览", font=("微软雅黑", 10, "bold"),
                 bg="white", fg="#2d5a8a").pack(anchor="w", padx=12, pady=(10, 4))
        preview_frame = tk.Frame(card_m, bg="#e8e0d8")
        preview_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        self.preview_canvas = tk.Canvas(preview_frame, bg="#e8e0d8",
                                        highlightthickness=0, scrollregion=(0, 0, 400, 400))
        self.preview_canvas.pack(fill="both", expand=True)
        self.preview_canvas.bind("<Configure>", self._on_preview_configure)

        # ── 右侧：设置（固定宽度 300） ──
        right = tk.Frame(body, width=300, bg="#f0f0f0")
        right.pack(side="right", fill="both", padx=(2, 4), pady=0)
        right.pack_propagate(False)
        ri = tk.Frame(right, bg="white", bd=1, relief="solid")
        ri.grid(row=0, column=0, sticky="nsew", padx=8, pady=8)
        ri.columnconfigure(0, weight=0, minsize=130)
        ri.columnconfigure(1, weight=1)
        right.columnconfigure(0, weight=1)

        r = [0]
        def R():
            r[0] += 1; return r[0]
        def title(text, sz=11):
            tk.Label(ri, text=text, font=("微软雅黑", sz, "bold"),
                     bg="white", fg="#2d5a8a", anchor="w").grid(
                         row=R(), column=0, columnspan=2, sticky="w", padx=14, pady=(12 if sz==11 else 4, 2))
        def sep():
            tk.Frame(ri, height=1, bg="#ddd").grid(
                row=R(), column=0, columnspan=2, sticky="ew", padx=14, pady=(5, 5))
        def row2(label, widget):
            tk.Label(ri, text=label, font=("微软雅黑", 9), bg="white", anchor="w").grid(
                row=R(), column=0, sticky="w", padx=(14, 6), pady=5)
            widget.grid(row=r[0], column=1, sticky="w", padx=(0, 14), pady=5)

        title("排版设置", 11)
        title_entry = ttk.Entry(ri, textvariable=self.doc_title, width=18, font=("微软雅黑", 9))
        row2("表格标题：", title_entry)
        title_entry.bind("<KeyRelease>", lambda e: self._upd_title_example())

        self.page_orient = tk.StringVar(value="纵")
        self.img_per_page = tk.IntVar(value=4)

        tk.Label(ri, text="纸张方向：", font=("微软雅黑", 9), bg="white", anchor="w").grid(
            row=R(), column=0, sticky="w", padx=(14, 6), pady=5)
        orient_frame = tk.Frame(ri, bg="white"); orient_frame.grid(row=r[0], column=1, sticky="w", pady=5)
        for val, lbl in [("纵", "A4纵版"), ("横", "A4横版")]:
            ttk.Radiobutton(orient_frame, text=lbl, variable=self.page_orient, value=val).pack(side="left", padx=6)

        tk.Label(ri, text="每页图片数：", font=("微软雅黑", 9), bg="white", anchor="w").grid(
            row=R(), column=0, sticky="w", padx=(14, 6), pady=5)
        page_frame = tk.Frame(ri, bg="white"); page_frame.grid(row=r[0], column=1, sticky="w", pady=5)
        for val, lbl in [(1, "1张"), (2, "2张"), (4, "4张"), (6, "6张"), (8, "8张")]:
            ttk.Radiobutton(page_frame, text=lbl, variable=self.img_per_page, value=val).pack(side="left", padx=4)

        tk.Checkbutton(ri, text="显示表格边框", variable=self.show_border,
                       font=("微软雅黑", 9), bg="white", anchor="w").grid(
                           row=R(), column=0, columnspan=2, sticky="w", padx=14, pady=5)
        sep()

        title("标注文字", 10)
        cap_entry = ttk.Entry(ri, textvariable=self.caption_tpl, width=18, font=("微软雅黑", 9))
        row2("标注模板：", cap_entry)
        tk.Label(ri, text="{n}=序号  {f}=文件名（不含扩展名）",
                 font=("微软雅黑", 7), fg="#888", bg="white").grid(
                     row=R(), column=0, columnspan=2, sticky="w", padx=20)
        s_spin = ttk.Spinbox(ri, from_=8, to=20, textvariable=self.caption_size, width=9)
        row2("标注字号：", s_spin)

        tk.Label(ri, text="标注对齐：", font=("微软雅黑", 9), bg="white", anchor="w").grid(
            row=R(), column=0, sticky="w", padx=(14, 6), pady=5)
        ab = tk.Frame(ri, bg="white"); ab.grid(row=r[0], column=1, sticky="w", pady=5)
        for val, lbl in [("居中", "居中"), ("左对齐", "左"), ("右对齐", "右")]:
            ttk.Radiobutton(ab, text=lbl, variable=self.caption_align, value=val).pack(side="left", padx=3)

        self.ex_label = tk.Label(ri, text="", font=("微软雅黑", 8),
                                  bg="#f3f3ff", fg="#3333aa", anchor="w", padx=6, pady=4)
        self.ex_label.grid(row=R(), column=0, columnspan=2, sticky="ew", padx=14, pady=(6, 2))
        cap_entry.bind("<KeyRelease>", lambda e: self._upd_example())

        sep()
        out_entry = ttk.Entry(ri, textvariable=self.output_name, width=18, font=("微软雅黑", 9))
        row2("Word 文件名：", out_entry)

        gen_btn = tk.Button(ri, text="生成 Word 文档",
                            command=self._generate, font=("微软雅黑", 11, "bold"),
                            bg="#217346", fg="white", relief="flat", cursor="hand2", height=2)
        gen_btn.grid(row=R(), column=0, columnspan=2, sticky="ew", padx=14, pady=(8, 4))

        self.status = tk.Label(self.root, text="就绪" + ("（拖拽可用）" if HAS_DND else "（点击添加）"), bd=1, relief="sunken",
                               anchor="w", font=("微软雅黑", 8), bg="#e8e8e8")
        self.status.pack(side="bottom", fill="x")

        self._upd_example()

    # ── 标题预览 ──────────────────────────────────────────────
    def _upd_title_example(self):
        t = self.doc_title.get().strip() or "估价对象现状照片"

    # ── 拖拽区域 ──────────────────────────────────────────────
    def _setup_drop_hint(self):
        self.drop_label.config(text="拖拽图片到此处", cursor="hand2")
        self.drop_frame.bind('<Button-1>', lambda e: self._add_images())
        if HAS_DND:
            self.drop_frame.drop_target_register(DND_FILES)
            self.drop_frame.dnd_bind('<<Drop>>', self._on_drop)

    def _on_drop(self, event):
        """处理拖拽放入的文件（文件夹/文件混合）。"""
        if not hasattr(event, 'data') or not event.data:
            return
        # tkinterdnd2 在 Windows 上返回的是 {file path} 格式或纯路径
        raw = event.data
        # 去掉首尾花括号（Windows URI 格式）
        raw = raw.strip()
        if raw.startswith('{') and raw.endswith('}'):
            raw = raw[1:-1]
        # 分割多文件（可能被空格分隔，但路径本身含空格时用\n或\t）
        parts = []
        # 尝试按换行、Tab 分隔
        for sep in ('\n', '\t'):
            if sep in raw:
                parts = raw.split(sep)
                break
        else:
            # 尝试按空格分割（不太准，但作为 fallback）
            parts = raw.split(' ')
        paths = []
        for p in parts:
            p = p.strip()
            if not p:
                continue
            # 去掉可能的引号
            p = p.strip('"').strip("'")
            if os.path.isdir(p):
                # 递归收集目录下所有图片
                for root_dir, _, files in os.walk(p):
                    for fn in sorted(files):
                        if self._is_image(fn):
                            paths.append(os.path.join(root_dir, fn))
            elif os.path.isfile(p):
                if self._is_image(p):
                    paths.append(p)
        if paths:
            self._add_files(paths)

    @staticmethod
    def _is_image(filename):
        return filename.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp'))

    # ── 功能 ─────────────────────────────────────────────────
    def _init_preview(self):
        self.root.after(50, self._do_preview)

    def _do_preview(self):
        """实际执行预览绘制，始终使用 self.current_idx"""
        cw = self.preview_canvas.winfo_width()
        ch = self.preview_canvas.winfo_height()
        if cw <= 1 or ch <= 1:
            self.root.after(100, self._do_preview)
            return
        self.preview_canvas.configure(scrollregion=(0, 0, cw, ch))
        idx = self.current_idx
        self.preview_canvas.delete("all")
        if idx < 0 or idx >= len(self.image_paths):
            cw0 = self.preview_canvas.winfo_width()
            ch0 = self.preview_canvas.winfo_height()
            self.preview_canvas.create_text(cw0 // 2, ch0 // 2,
                                            text="← 从左侧列表选择图片预览",
                                            font=("微软雅黑", 12), fill="#bbb", anchor="center")
            return
        try:
            path = self.image_paths[idx]
            pil_img = Image.open(path)
            max_w = max(cw - 20, 50)
            max_h = max(ch - 20, 50)
            pil_img.thumbnail((max_w, max_h), Image.LANCZOS)
            photo = ImageTk.PhotoImage(pil_img)
            self._preview_img = photo
            x = cw // 2
            y = ch // 2
            self.preview_canvas.create_image(x, y, anchor="center", image=photo)
            self.preview_canvas.create_text(x, y + pil_img.height // 2 + 12,
                                            text=os.path.basename(path),
                                            font=("微软雅黑", 8), fill="#666", anchor="n")
            self.preview_canvas.create_text(x, y + pil_img.height // 2 + 28,
                                            text=f"{pil_img.width} × {pil_img.height}",
                                            font=("微软雅黑", 7), fill="#aaa", anchor="n")
        except Exception as e:
            self.preview_canvas.delete("all")
            self.preview_canvas.create_text(cw // 2, ch // 2,
                                            text="预览失败:\n" + str(e),
                                            font=("微软雅黑", 9), fill="#f88", anchor="center")

    def _on_preview_configure(self, event):
        self.root.after(50, self._do_preview)

    def _add_images(self):
        files = filedialog.askopenfilenames(title="选择图片",
            filetypes=[("图片", "*.jpg *.jpeg *.png *.bmp *.gif *.tiff *.webp"), ("所有文件", "*.*")])
        if files:
            self._add_files(files)

    def _add_files(self, files):
        added = 0
        for f in files:
            if f not in self.image_paths:
                self.image_paths.append(f)
                self.listbox.insert(tk.END, os.path.basename(f))
                added += 1
        if added > 0:
            self._upd_status()
            self.listbox.selection_clear(0, tk.END)
            self.listbox.selection_set(tk.END)
            self.listbox.see(tk.END)
            self._on_select(None)

    def _remove_sel(self):
        sel = sorted(self.listbox.curselection(), reverse=True)
        for i in sel:
            self.listbox.delete(i); del self.image_paths[i]
        self.current_idx = -1
        self.root.after_idle(self._do_preview)
        self._upd_status()

    def _shift(self, direction):
        sel = list(self.listbox.curselection())
        if not sel: return
        idx = sel[0]; ni = idx + direction
        if 0 <= ni < len(self.image_paths):
            self.image_paths[idx], self.image_paths[ni] = self.image_paths[ni], self.image_paths[idx]
            self.listbox.delete(0, tk.END)
            for p in self.image_paths: self.listbox.insert(tk.END, os.path.basename(p))
            self.listbox.selection_set(ni); self.listbox.see(ni)
            self.current_idx = ni
            self.root.after_idle(self._do_preview)

    def _clear_all(self):
        if not self.image_paths: return
        if not messagebox.askyesno("确认", "确定清空所有图片？"): return
        self.listbox.delete(0, tk.END); self.image_paths.clear()
        self.current_idx = -1
        self.root.after_idle(self._do_preview)
        self._upd_status()

    def _open_img(self):
        sel = list(self.listbox.curselection())
        if sel:
            try: os.startfile(self.image_paths[sel[0]])
            except Exception: pass

    def _on_select(self, event):
        sel = list(self.listbox.curselection())
        if sel:
            self.current_idx = sel[-1]
        else:
            self.current_idx = -1
        self.root.after_idle(self._do_preview)

    def _upd_example(self):
        tpl = self.caption_tpl.get()
        if self.image_paths:
            fname = os.path.splitext(os.path.basename(self.image_paths[0]))[0]
            ex = tpl.replace("{n}", "1").replace("{f}", fname) if tpl else "(无标注)"
        else:
            ex = tpl.replace("{n}", "1").replace("{f}", "示例图片") if tpl else "(无标注)"
        self.ex_label.config(text="  预览：" + ex)

    def _upd_status(self):
        n = len(self.image_paths)
        base = "共 " + str(n) + " 张图片" if n else "就绪"
        self.status.config(text=base + ("（拖拽可用）" if HAS_DND else "（点击添加）"))

    def _generate(self):
        if not self.image_paths:
            messagebox.showwarning("提示", "请先添加图片！"); return

        out = filedialog.asksaveasfilename(title="保存 Word",
            defaultextension=".docx",
            initialfile=self.output_name.get().strip() or "图片排版文档",
            filetypes=[("Word 文档", "*.docx")])
        if not out: return

        self.status.config(text="正在生成...")
        self.root.update_idletasks()

        try:
            from PIL import Image as PILImage
            doc = Document()
            sec = doc.sections[0]
            cap_tpl = self.caption_tpl.get()
            cap_pt  = self.caption_size.get()
            rgb = (0x33, 0x33, 0x33)  # 固定深灰色
            align_map = {"居中": WD_ALIGN_PARAGRAPH.CENTER,
                         "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
                         "右对齐": WD_ALIGN_PARAGRAPH.RIGHT}
            cap_align = align_map.get(self.caption_align.get(), WD_ALIGN_PARAGRAPH.CENTER)

            # 纸张方向与布局
            is_landscape = (self.page_orient.get() == "横")
            images_per_page = self.img_per_page.get()
            show_border = self.show_border.get()

            # 直接操作底层 XML 设置页面尺寸和边距
            pg_sz = sec._sectPr.find(qn('w:pgSz'))
            if is_landscape:
                pg_sz.set(qn('w:w'), str(int(29.7 / 2.54 * 1440)))
                pg_sz.set(qn('w:h'), str(int(21 / 2.54 * 1440)))
                pg_sz.set(qn('w:orient'), 'landscape')
            else:
                pg_sz.set(qn('w:w'), str(int(21 / 2.54 * 1440)))
                pg_sz.set(qn('w:h'), str(int(29.7 / 2.54 * 1440)))
                pg_sz.set(qn('w:orient'), 'portrait')
            pg_mar = sec._sectPr.find(qn('w:pgMar'))
            if pg_mar is None:
                pg_mar = OxmlElement('w:pgMar')
                sec._sectPr.append(pg_mar)
            for attr, val in [('w:top', 1.1), ('w:bottom', 1.2), ('w:left', 2.4), ('w:right', 1.7)]:
                pg_mar.set(qn(attr), str(int(val / 2.54 * 1440)))

            # 标题（仅当有内容时添加）
            title_text = self.doc_title.get().strip()
            TITLE_H_ACTUAL = 1.2 if title_text else 0.0
            if title_text:
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tr = title_para.add_run(title_text)
                tr.font.size = Pt(12)
                tr.font.bold = True
                tr.font.name = '黑体'
                tr._r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            # ── 预读所有图片宽高比 ──────────────────────────────────
            n = len(self.image_paths)
            aspects = []   # height/width for each image
            for p in self.image_paths:
                try:
                    with PILImage.open(p) as im:
                        w, h = im.size
                        aspects.append(h / w if w > 0 else 1.0)
                except Exception:
                    aspects.append(1.0)

            # 可用页面尺寸（厘米）
            PAGE_W, PAGE_H = 21.0, 29.7
            MARGIN_L, MARGIN_R, MARGIN_T, MARGIN_B = 2.4, 1.7, 1.1, 1.2
            LABEL_H = 1.0

            avail_w = PAGE_W - MARGIN_L - MARGIN_R
            avail_h = PAGE_H - MARGIN_T - MARGIN_B - TITLE_H_ACTUAL

            # 每行2列，1张/页时用1列（单图满宽）特殊处理
            n_pairs = (n + 1) // 2
            if images_per_page == 1:
                rows_per_page = n_pairs  # 1行1对（1图）
            else:
                rows_per_page = images_per_page // 2
            row_h_cm = avail_h / rows_per_page
            col_w_in = avail_w / 2 / 2.54  # 每格宽度（英寸）

            single_w_in = avail_w / 2.54   # 单列全宽（1张/页时用）

            # 按行分组处理
            if images_per_page == 1:
                # 1张/页：每张图单独一行（满宽）
                img_idx = 0
                while img_idx < n:
                    img_row = doc.add_table(rows=1, cols=1)
                    img_row.style = 'Table Grid'
                    img_row.alignment = 1
                    set_cell_width(img_row.rows[0].cells[0], single_w_in)
                    add_cell_image(img_row.rows[0].cells[0], self.image_paths[img_idx], single_w_in)
                    if not show_border:
                        remove_table_borders(img_row)

                    if cap_tpl.strip():
                        cap_row = doc.add_table(rows=1, cols=1)
                        cap_row.style = 'Table Grid'
                        cap_row.alignment = 1
                        set_cell_width(cap_row.rows[0].cells[0], single_w_in)
                        fname = os.path.splitext(os.path.basename(self.image_paths[img_idx]))[0]
                        add_cell_text(cap_row.rows[0].cells[0],
                                      cap_tpl.replace("{n}", str(img_idx + 1)).replace("{f}", fname),
                                      cap_pt, rgb, cap_align)
                        if not show_border:
                            remove_table_borders(cap_row)
                    img_idx += 1
            else:
                pair_idx = 0
                while pair_idx < n_pairs:
                    rows_in_this_page = min(rows_per_page, n_pairs - pair_idx)
                    for _ in range(rows_in_this_page):
                        if pair_idx >= n_pairs:
                            break
                        left_i  = pair_idx * 2
                        right_i = pair_idx * 2 + 1
                        has_right = right_i < n

                        img_row = doc.add_table(rows=1, cols=2)
                        img_row.style = 'Table Grid'
                        img_row.alignment = 1
                        set_cell_width(img_row.rows[0].cells[0], col_w_in)
                        set_cell_width(img_row.rows[0].cells[1], col_w_in)
                        add_cell_image(img_row.rows[0].cells[0], self.image_paths[left_i], col_w_in)
                        if has_right:
                            add_cell_image(img_row.rows[0].cells[1], self.image_paths[right_i], col_w_in)
                        if not show_border:
                            remove_table_borders(img_row)

                        if cap_tpl.strip():
                            cap_row = doc.add_table(rows=1, cols=2)
                            cap_row.style = 'Table Grid'
                            cap_row.alignment = 1
                            set_cell_width(cap_row.rows[0].cells[0], col_w_in)
                            set_cell_width(cap_row.rows[0].cells[1], col_w_in)
                            left_fname = os.path.splitext(os.path.basename(self.image_paths[left_i]))[0]
                            add_cell_text(cap_row.rows[0].cells[0],
                                          cap_tpl.replace("{n}", str(left_i + 1)).replace("{f}", left_fname),
                                          cap_pt, rgb, cap_align)
                            if has_right:
                                right_fname = os.path.splitext(os.path.basename(self.image_paths[right_i]))[0]
                                add_cell_text(cap_row.rows[0].cells[1],
                                              cap_tpl.replace("{n}", str(right_i + 1)).replace("{f}", right_fname),
                                              cap_pt, rgb, cap_align)
                            else:
                                cap_row.rows[0].cells[1].paragraphs[0].clear()
                            if not show_border:
                                remove_table_borders(cap_row)
                        pair_idx += 1

            doc.save(out)
            self.status.config(text="已生成：" + os.path.basename(out))
            messagebox.showinfo("完成", "Word 文档已生成！\n\n" + out)
            os.startfile(out)

        except Exception as e:
            import traceback, io, sys
            buf = io.StringIO()
            traceback.print_exc(file=buf)
            tb_text = buf.getvalue()
            self.status.config(text="生成失败")
            # 写入日志文件
            log_path = os.path.join(os.path.dirname(__file__), 'run_log.txt')
            with open(log_path, 'w', encoding='utf-8') as lf:
                lf.write("=== GENERATE ERROR ===\n")
                lf.write(tb_text)
                lf.write("\n======================\n")
            print("=== GENERATE ERROR ===", file=sys.stdout, flush=True)
            print(tb_text, file=sys.stdout, flush=True)
            print("======================", file=sys.stdout, flush=True)
            messagebox.showerror("错误", "生成失败：\n" + str(e) + "\n\n详情已打印到控制台")


if __name__ == "__main__":
    if HAS_DND:
        root = TkinterDnD.Tk()
    else:
        root = tk.Tk()
    app = ImageToWordApp(root)
    root.mainloop()
