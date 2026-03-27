"""测试 JPEG 输出"""
import sys, os, tempfile
sys.stdout.reconfigure(encoding='utf-8')
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches

doc = Document()
p = doc.add_paragraph()
run = p.add_run()
img_path = r'C:\Users\Administrator\Desktop\新建文件夹\01储油罐.jpg'
print(f"Source: {img_path}")

with PILImage.open(img_path) as pil_img:
    print(f"Mode: {pil_img.mode}, Size: {pil_img.size}")
    if pil_img.mode in ('P','LA','PA','CMYK','RGBA'):
        clean = pil_img.convert('RGBA')
        bg = PILImage.new('RGBA', clean.size, (255,255,255,255))
        bg.alpha_composite(clean)
        clean = bg.convert('RGB')
    elif pil_img.mode != 'RGB':
        clean = pil_img.convert('RGB')
    else:
        clean = pil_img
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.jpg')
    os.close(tmp_fd)
    clean.save(tmp_path, format='JPEG', quality=95)
    size = os.path.getsize(tmp_path)
    print(f"Temp JPEG: {size:,} bytes")

run.add_picture(tmp_path, width=Inches(3.0))
os.remove(tmp_path)

doc.save(os.path.join(os.environ['TEMP'], 'jpeg_test.docx'))
print("Saved OK")
