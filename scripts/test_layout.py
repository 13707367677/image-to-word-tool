"""Test 4-per-page auto-scaling layout"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, r'C:\Users\Administrator\lobsterai\project\scripts')

from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile

def set_cell_width(cell, width_in):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(int(width_in * 1440)))
    tc_w.set(qn('w:type'), 'dxa')
    tc_pr.append(tc_w)

def add_cell_image(cell, img_path, width_in):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    with PILImage.open(img_path) as pil_img:
        if pil_img.mode in ('P','LA','PA','CMYK','RGBA'):
            clean = pil_img.convert('RGBA')
            bg = PILImage.new('RGBA', clean.size, (255,255,255,255))
            bg.alpha_composite(clean); clean = bg.convert('RGB')
        elif pil_img.mode != 'RGB':
            clean = pil_img.convert('RGB')
        else:
            clean = pil_img
        fd, path = tempfile.mkstemp(suffix='.jpg')
        os.close(fd)
        clean.save(path, format='JPEG', quality=95)
        run.add_picture(path, width=Inches(width_in))
        os.remove(path)

def add_cell_text(cell, text, size_pt, rgb, align):
    p = cell.paragraphs[0]; p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*rgb)

from docx.shared import RGBColor

def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr'); tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),'none'); b.set(qn('w:sz'),'0')
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),'auto')
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)

photo_dir = r'C:\Users\Administrator\Desktop\新建文件夹'
photos = sorted([os.path.join(photo_dir, f) for f in os.listdir(photo_dir)
                 if f.lower().endswith(('.jpg','.jpeg','.png'))])
print(f"Photos: {len(photos)}")

for images_per_page in [2, 4, 6, 8]:
    doc = Document()
    sec = doc.sections[0]
    pg_sz = sec._sectPr.find(qn('w:pgSz'))
    pg_sz.set(qn('w:w'), str(int(21 / 2.54 * 1440)))
    pg_sz.set(qn('w:h'), str(int(29.7 / 2.54 * 1440)))
    pg_sz.set(qn('w:orient'), 'portrait')
    pg_mar = sec._sectPr.find(qn('w:pgMar'))
    if pg_mar is None:
        pg_mar = OxmlElement('w:pgMar'); sec._sectPr.append(pg_mar)
    for attr, val in [('w:top',1.1),('w:bottom',1.2),('w:left',2.4),('w:right',1.7)]:
        pg_mar.set(qn(attr), str(int(val / 2.54 * 1440)))

    PAGE_W, PAGE_H = 21.0, 29.7
    avail_w = PAGE_W - 2.4 - 1.7  # 16.9cm
    avail_h = PAGE_H - 1.1 - 1.2 - 1.2  # 26.2cm

    aspects = []
    for p in photos:
        with PILImage.open(p) as im:
            w, h = im.size
            aspects.append(h/w if w > 0 else 1.0)

    n = len(photos)
    n_pairs = (n + 1) // 2
    rows_per_page = images_per_page // 2
    row_h_cm = avail_h / rows_per_page
    col_w_in = avail_w / 2 / 2.54

    cap_tpl = "{f}"
    rgb = (0x33, 0x33, 0x33)
    cap_align = WD_ALIGN_PARAGRAPH.CENTER

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("估价对象现状照片")
    tr.font.size = Pt(12); tr.font.bold = True

    pair_idx = 0
    while pair_idx < n_pairs:
        rows_in_page = min(rows_per_page, n_pairs - pair_idx)
        for _ in range(rows_in_page):
            if pair_idx >= n_pairs:
                break
            left_i = pair_idx * 2
            right_i = pair_idx * 2 + 1
            has_right = right_i < n

            img_row = doc.add_table(rows=1, cols=2)
            img_row.style = 'Table Grid'; img_row.alignment = 1
            set_cell_width(img_row.rows[0].cells[0], col_w_in)
            set_cell_width(img_row.rows[0].cells[1], col_w_in)
            add_cell_image(img_row.rows[0].cells[0], photos[left_i], col_w_in)
            if has_right:
                add_cell_image(img_row.rows[0].cells[1], photos[right_i], col_w_in)
            remove_table_borders(img_row)

            cap_row = doc.add_table(rows=1, cols=2)
            cap_row.style = 'Table Grid'; cap_row.alignment = 1
            set_cell_width(cap_row.rows[0].cells[0], col_w_in)
            set_cell_width(cap_row.rows[0].cells[1], col_w_in)
            lf = os.path.splitext(os.path.basename(photos[left_i]))[0]
            add_cell_text(cap_row.rows[0].cells[0], cap_tpl.replace("{n}",str(left_i+1)).replace("{f}",lf), 10, rgb, cap_align)
            if has_right:
                rf = os.path.splitext(os.path.basename(photos[right_i]))[0]
                add_cell_text(cap_row.rows[0].cells[1], cap_tpl.replace("{n}",str(right_i+1)).replace("{f}",rf), 10, rgb, cap_align)
            remove_table_borders(cap_row)
            pair_idx += 1

    out = os.path.join(os.environ['TEMP'], f'test_4perpage_{images_per_page}.docx')
    doc.save(out)
    print(f"{images_per_page}张/页: {os.path.getsize(out):,} bytes")

print("SUCCESS")
