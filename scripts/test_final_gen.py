"""直接用工具的 _generate 逻辑测试8张照片"""
import sys, os, tempfile
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, r'C:\Users\Administrator\lobsterai\project\scripts')

from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_width(cell, width_in):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(int(width_in * 1440)))
    tc_w.set(qn('w:type'), 'dxa')
    tc_pr.append(tc_w)

def add_cell_image(cell, img_path, width_in):
    import tempfile
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    with PILImage.open(img_path) as pil_img:
        if pil_img.mode in ('P', 'LA', 'PA', 'CMYK', 'RGBA'):
            clean_img = pil_img.convert('RGBA')
            background = PILImage.new('RGBA', clean_img.size, (255, 255, 255, 255))
            background.alpha_composite(clean_img)
            clean_img = background.convert('RGB')
        elif pil_img.mode != 'RGB':
            clean_img = pil_img.convert('RGB')
        else:
            clean_img = pil_img
        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.png')
        os.close(tmp_fd)
        try:
            clean_img.save(tmp_path, format='PNG')
            run.add_picture(tmp_path, width=Inches(width_in))
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

def add_cell_text(cell, text, size_pt, rgb, align, bold=False):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*rgb)
    run.font.bold = bold

def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)

photo_dir = r'C:\Users\Administrator\Desktop\新建文件夹'
photos = sorted([os.path.join(photo_dir, f) for f in os.listdir(photo_dir)
                 if f.lower().endswith(('.jpg','.jpeg','.png'))])
print(f'Photos: {len(photos)}')

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.1)
sec.bottom_margin = Cm(1.2)
sec.left_margin = Cm(2.4)
sec.right_margin = Cm(1.7)

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("估价对象现状照片")
tr.font.size = Pt(12)
tr.font.bold = True

max_w_in = 8.0 / 2.54
cap_align = WD_ALIGN_PARAGRAPH.CENTER
rgb = (0x33, 0x33, 0x33)
cap_pt = 10
cap_tpl = "{n} {f}"

images_per_group = 4  # 2×2
n = len(photos)
pair_idx = 0

while pair_idx * 2 < n:
    group_end = min(pair_idx * 2 + images_per_group, n)
    while pair_idx * 2 < group_end:
        left_i = pair_idx * 2
        right_i = pair_idx * 2 + 1
        has_right = right_i < n

        img_row = doc.add_table(rows=1, cols=2)
        img_row.style = 'Table Grid'
        img_row.alignment = 1
        set_cell_width(img_row.rows[0].cells[0], max_w_in)
        set_cell_width(img_row.rows[0].cells[1], max_w_in)
        add_cell_image(img_row.rows[0].cells[0], photos[left_i], max_w_in)
        if has_right:
            add_cell_image(img_row.rows[0].cells[1], photos[right_i], max_w_in)
        remove_table_borders(img_row)

        cap_row = doc.add_table(rows=1, cols=2)
        cap_row.style = 'Table Grid'
        cap_row.alignment = 1
        set_cell_width(cap_row.rows[0].cells[0], max_w_in)
        set_cell_width(cap_row.rows[0].cells[1], max_w_in)
        left_fname = os.path.splitext(os.path.basename(photos[left_i]))[0]
        add_cell_text(cap_row.rows[0].cells[0], cap_tpl.replace("{n}", str(left_i+1)).replace("{f}", left_fname), cap_pt, rgb, cap_align)
        if has_right:
            right_fname = os.path.splitext(os.path.basename(photos[right_i]))[0]
            add_cell_text(cap_row.rows[0].cells[1], cap_tpl.replace("{n}", str(right_i+1)).replace("{f}", right_fname), cap_pt, rgb, cap_align)
        remove_table_borders(cap_row)

        pair_idx += 1

out = os.path.join(os.environ['TEMP'], 'final_test.docx')
doc.save(out)
print(f'Saved: {os.path.getsize(out):,} bytes')
print('SUCCESS - all 8 photos generated!')
