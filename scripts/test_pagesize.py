"""Test page size via XML"""
from docx import Document
from docx.oxml.ns import qn

doc = Document()
sec = doc.sections[0]

# Set via XML directly
pg_sz = sec._sectPr.find(qn('w:pgSz'))
print("pgSz before:", pg_sz.attrib if pg_sz is not None else "None")

# A4: 21cm x 29.7cm, width > height means landscape
w_twips = int(29.7 / 2.54 * 1440)
h_twips = int(21 / 2.54 * 1440)
pg_sz.set(qn('w:w'), str(w_twips))
pg_sz.set(qn('w:h'), str(h_twips))
pg_sz.set(qn('w:orient'), 'landscape')

# Set margins via XML
pg_mar = sec._sectPr.find(qn('w:pgMar'))
if pg_mar is None:
    pg_mar = __import__('xml.etree.ElementTree').SubElement(sec._sectPr, qn('w:pgMar'))
top = int(1.1 / 2.54 * 1440)
bot = int(1.2 / 2.54 * 1440)
lft = int(2.4 / 2.54 * 1440)
rgt = int(1.7 / 2.54 * 1440)
pg_mar.set(qn('w:top'), str(top))
pg_mar.set(qn('w:bottom'), str(bot))
pg_mar.set(qn('w:left'), str(lft))
pg_mar.set(qn('w:right'), str(rgt))

doc.add_paragraph("A4横向测试")

doc.save("C:/Users/ADMINI~1/AppData/Local/Temp/a4_landscape.docx")

# Verify
doc2 = Document("C:/Users/ADMINI~1/AppData/Local/Temp/a4_landscape.docx")
sec2 = doc2.sections[0]
pg_sz2 = sec2._sectPr.find(qn('w:pgSz'))
w2 = int(pg_sz2.get(qn('w:w'))) / 1440 * 2.54
h2 = int(pg_sz2.get(qn('w:h'))) / 1440 * 2.54
pg_mar2 = sec2._sectPr.find(qn('w:pgMar'))
t2 = int(pg_mar2.get(qn('w:top'))) / 1440 * 2.54
b2 = int(pg_mar2.get(qn('w:bottom'))) / 1440 * 2.54
l2 = int(pg_mar2.get(qn('w:left'))) / 1440 * 2.54
r2 = int(pg_mar2.get(qn('w:right'))) / 1440 * 2.54
print(f"Loaded: {w2:.1f}cm x {h2:.1f}cm  top={t2:.2f} bot={b2:.2f} left={l2:.2f} right={r2:.2f}")
