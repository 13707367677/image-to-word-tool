import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from PIL import Image as PILImage
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_width(cell, width_in):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(int(width_in * 1440)))
    tc_w.set(qn('w:type'), 'dxa')
    tc_pr.append(tc_w)

def add_cell_image(cell, img_path, width_in):
    from PIL import Image as PILImage
    from io import BytesIO
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    buf = BytesIO()
    with PILImage.open(img_path) as pil_img:
        pil_img.save(buf, format=pil_img.format or 'JPEG', quality=95)
    buf.seek(0)
    run.add_picture(buf, width=Inches(width_in))

def add_cell_text(cell, text, size_pt, rgb, align):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*rgb)

# Test with actual user image from the template dir
import glob
test_dir = r'C:\Users\ADMINI~1\AppData\Local\Temp\lobsterai\attachments'
imgs = glob.glob(test_dir + '/*.jpg') + glob.glob(test_dir + '/*.jpeg')
print('Found images:', imgs[:3])

images = imgs[:4] if imgs else []
if not images:
    # Use fallback
    images = [r'C:\Users\Administrator\lobsterai\project\after_login.png']
    print('Using fallback:', images)

cap_tpl = "{n} {f}"
cap_pt = 10
rgb = (51, 51, 51)
cap_align = WD_ALIGN_PARAGRAPH.CENTER
max_w_in = 8.0 / 2.54

out = os.path.join(os.environ['TEMP'], 'test_fixed.docx')

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.1)
sec.bottom_margin = Cm(1.2)
sec.left_margin = Cm(2.4)
sec.right_margin = Cm(1.7)

n = len(images)
pair_idx = 0

while pair_idx * 2 < n:
    left_i  = pair_idx * 2
    right_i = pair_idx * 2 + 1
    has_right = right_i < n

    print(f'Pair {pair_idx}: left={left_i} right={right_i} has_right={has_right}')
    print('  Left image:', images[left_i])

    img_row = doc.add_table(rows=1, cols=2)
    img_row.style = 'Table Grid'
    set_cell_width(img_row.rows[0].cells[0], max_w_in)
    set_cell_width(img_row.rows[0].cells[1], max_w_in)
    add_cell_image(img_row.rows[0].cells[0], images[left_i], max_w_in)
    print('  Left OK')
    if has_right:
        print('  Right image:', images[right_i])
        add_cell_image(img_row.rows[0].cells[1], images[right_i], max_w_in)
        print('  Right OK')

    cap_row = doc.add_table(rows=1, cols=2)
    cap_row.style = 'Table Grid'
    set_cell_width(cap_row.rows[0].cells[0], max_w_in)
    set_cell_width(cap_row.rows[0].cells[1], max_w_in)
    left_fname = os.path.splitext(os.path.basename(images[left_i]))[0]
    left_cap = cap_tpl.replace("{n}", str(left_i + 1)).replace("{f}", left_fname)
    add_cell_text(cap_row.rows[0].cells[0], left_cap, cap_pt, rgb, cap_align)
    if has_right:
        right_fname = os.path.splitext(os.path.basename(images[right_i]))[0]
        right_cap = cap_tpl.replace("{n}", str(right_i + 1)).replace("{f}", right_fname)
        add_cell_text(cap_row.rows[0].cells[1], right_cap, cap_pt, rgb, cap_align)

    pair_idx += 1

doc.save(out)
print('Saved to:', out)
print('File size:', os.path.getsize(out), 'bytes')
print('SUCCESS')
