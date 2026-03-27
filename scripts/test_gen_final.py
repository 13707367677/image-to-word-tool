import sys, os
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, r'C:\Users\Administrator\lobsterai\project\scripts')

from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_width(cell, width_in):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(int(width_in * 1440)))
    tc_w.set(qn('w:type'), 'dxa')
    tc_pr.append(tc_w)

def add_cell_image(cell, img_path, width_in):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_in))

def add_cell_text(cell, text, size_pt, rgb, align):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*rgb)

# Exactly replicate the _generate function
images = [r'C:\Users\Administrator\lobsterai\project\after_login.png']
cap_tpl = "{n} {f}"
cap_pt = 10
rgb = (51, 51, 51)
cap_align = WD_ALIGN_PARAGRAPH.CENTER
max_w_in = 8.0 / 2.54

out = r'C:\Users\Administrator\lobsterai\project\test_output.docx'

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.1)
sec.bottom_margin = Cm(1.2)
sec.left_margin = Cm(2.4)
sec.right_margin = Cm(1.7)

n = len(images)
pair_idx = 0

while pair_idx * 2 < n:
    left_i  = pair_idx * 2
    right_i = pair_idx * 2 + 1
    has_right = right_i < n

    # Image row
    img_row = doc.add_table(rows=1, cols=2)
    img_row.style = 'Table Grid'
    set_cell_width(img_row.rows[0].cells[0], max_w_in)
    set_cell_width(img_row.rows[0].cells[1], max_w_in)
    add_cell_image(img_row.rows[0].cells[0], images[left_i], max_w_in)
    if has_right:
        add_cell_image(img_row.rows[0].cells[1], images[right_i], max_w_in)

    # Caption row
    cap_row = doc.add_table(rows=1, cols=2)
    cap_row.style = 'Table Grid'
    set_cell_width(cap_row.rows[0].cells[0], max_w_in)
    set_cell_width(cap_row.rows[0].cells[1], max_w_in)

    left_fname = os.path.splitext(os.path.basename(images[left_i]))[0]
    left_cap = cap_tpl.replace("{n}", str(left_i + 1)).replace("{f}", left_fname)
    add_cell_text(cap_row.rows[0].cells[0], left_cap, cap_pt, rgb, cap_align)

    if has_right:
        right_fname = os.path.splitext(os.path.basename(images[right_i]))[0]
        right_cap = cap_tpl.replace("{n}", str(right_i + 1)).replace("{f}", right_fname)
        add_cell_text(cap_row.rows[0].cells[1], right_cap, cap_pt, rgb, cap_align)

    pair_idx += 1

doc.save(out)
print('Saved to:', out)
print('File size:', os.path.getsize(out), 'bytes')

# Verify
import zipfile
with zipfile.ZipFile(out) as z:
    imgs = [n for n in z.namelist() if 'media' in n]
    print('Embedded images:', imgs)
print('SUCCESS')
