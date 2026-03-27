"""直接测试用户真实图片"""
import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')

from PIL import Image as PILImage
from io import BytesIO

att_dir = r'C:\Users\ADMINI~1\AppData\Local\Temp\lobsterai\attachments'
imgs = glob.glob(att_dir + '/*.jpg') + glob.glob(att_dir + '/*.jpeg')
print('JPEG images found:', len(imgs))
if not imgs:
    imgs = glob.glob(att_dir + '/*')
    print('All files:', [os.path.basename(f) for f in imgs[:10]])

for img_path in imgs[:2]:
    print(f'\nTesting: {img_path}')
    try:
        with PILImage.open(img_path) as pil_img:
            print(f'  Mode: {pil_img.mode}, Size: {pil_img.size}')
            if pil_img.mode in ('P', 'LA', 'PA', 'CMYK', 'RGBA'):
                clean_img = pil_img.convert('RGBA')
                background = PILImage.new('RGBA', clean_img.size, (255, 255, 255, 255))
                background.alpha_composite(clean_img)
                clean_img = background.convert('RGB')
                print(f'  Converted to RGB: {clean_img.mode}, Size: {clean_img.size}')
            elif pil_img.mode != 'RGB':
                clean_img = pil_img.convert('RGB')
                print(f'  Converted to RGB: {clean_img.mode}')
            else:
                clean_img = pil_img

            buf = BytesIO()
            clean_img.save(buf, format='PNG')
            buf.seek(0)
            data = buf.read()
            print(f'  PNG buffer size: {len(data)} bytes')
            print(f'  PNG magic: {data[:8].hex()}')

            # Now test passing to docx
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = Document()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            buf.seek(0)
            run.add_picture(buf, width=Inches(3.0))
            out = os.path.join(os.environ['TEMP'], 'single_img_test.docx')
            doc.save(out)
            print(f'  docx saved OK: {os.path.getsize(out)} bytes')
    except Exception as e:
        print(f'  ERROR: {type(e).__name__}: {e}')
        import traceback; traceback.print_exc()

print('\nAll tests done')
