"""测试生成：纯PNG中转方案"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_width(cell, width_in):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(int(width_in * 1440)))
    tc_w.set(qn('w:type'), 'dxa')
    tc_pr.append(tc_w)

def add_cell_image(cell, img_path, width_in):
    """Add picture via pure-RGB PNG BytesIO to bypass EXIF/TIFF parsing bugs in python-docx."""
    from PIL import Image as PILImage
    from io import BytesIO
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    with PILImage.open(img_path) as pil_img:
        if pil_img.mode in ('P', 'LA', 'PA', 'CMYK', 'RGBA'):
            clean_img = pil_img.convert('RGBA')
            background = PILImage.new('RGBA', clean_img.size, (255, 255, 255, 255))
            background.alpha_composite(clean_img)
            clean_img = background.convert('RGB')
        elif pil_img.mode != 'RGB':
            clean_img = pil_img.convert('RGB')
        else:
            clean_img = pil_img
        buf = BytesIO()
        clean_img.save(buf, format='PNG')
    buf.seek(0)
    run.add_picture(buf, width=Inches(width_in))

def add_cell_text(cell, text, size_pt, rgb, align, bold=False):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*rgb)
    run.font.bold = bold

# 测试：用模板目录下的图片
import glob
att_dir = r'C:\Users\ADMINI~1\AppData\Local\Temp\lobsterai\attachments'
imgs = glob.glob(att_dir + '/*.jpg') + glob.glob(att_dir + '/*.jpeg') + glob.glob(att_dir + '/*.png')
print('Found images:', len(imgs))
if imgs:
    print('First:', imgs[0])
else:
    imgs = [r'C:\Users\Administrator\lobsterai\project\after_login.png']
    print('Using fallback:', imgs[0])

out = os.path.join(os.environ['TEMP'], 'png_test.docx')
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.1)
sec.bottom_margin = Cm(1.2)
sec.left_margin = Cm(2.4)
sec.right_margin = Cm(1.7)

# Title
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("估价对象现状照片")
tr.font.size = Pt(12)
tr.font.bold = True

max_w_in = 8.0 / 2.54
cap_align = WD_ALIGN_PARAGRAPH.CENTER
rgb = (0x33, 0x33, 0x33)
cap_pt = 10
cap_tpl = "{n} {f}"

n = len(imgs)
pair_idx = 0
while pair_idx * 2 < n:
    left_i = pair_idx * 2
    right_i = pair_idx * 2 + 1
    has_right = right_i < n

    img_row = doc.add_table(rows=1, cols=2)
    img_row.style = 'Table Grid'
    img_row.alignment = 1
    set_cell_width(img_row.rows[0].cells[0], max_w_in)
    set_cell_width(img_row.rows[0].cells[1], max_w_in)

    print(f'Adding left image {left_i}: {os.path.basename(imgs[left_i])}')
    add_cell_image(img_row.rows[0].cells[0], imgs[left_i], max_w_in)
    if has_right:
        print(f'Adding right image {right_i}: {os.path.basename(imgs[right_i])}')
        add_cell_image(img_row.rows[0].cells[1], imgs[right_i], max_w_in)

    cap_row = doc.add_table(rows=1, cols=2)
    cap_row.style = 'Table Grid'
    cap_row.alignment = 1
    set_cell_width(cap_row.rows[0].cells[0], max_w_in)
    set_cell_width(cap_row.rows[0].cells[1], max_w_in)

    left_fname = os.path.splitext(os.path.basename(imgs[left_i]))[0]
    add_cell_text(cap_row.rows[0].cells[0], cap_tpl.replace("{n}", str(left_i+1)).replace("{f}", left_fname), cap_pt, rgb, cap_align)
    if has_right:
        right_fname = os.path.splitext(os.path.basename(imgs[right_i]))[0]
        add_cell_text(cap_row.rows[0].cells[1], cap_tpl.replace("{n}", str(right_i+1)).replace("{f}", right_fname), cap_pt, rgb, cap_align)

    pair_idx += 1

doc.save(out)
print(f'Saved to: {out}')
print(f'File size: {os.path.getsize(out)} bytes')
print('SUCCESS')
