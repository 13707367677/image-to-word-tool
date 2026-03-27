"""Full generation test with correct page size via XML"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, r'C:\Users\Administrator\lobsterai\project\scripts')

from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile

def set_cell_width(cell, width_in):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), str(int(width_in * 1440)))
    tc_w.set(qn('w:type'), 'dxa')
    tc_pr.append(tc_w)

def add_cell_image(cell, img_path, width_in):
    import tempfile
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    with PILImage.open(img_path) as pil_img:
        if pil_img.mode in ('P', 'LA', 'PA', 'CMYK', 'RGBA'):
            clean = pil_img.convert('RGBA')
            bg = PILImage.new('RGBA', clean.size, (255,255,255,255))
            bg.alpha_composite(clean); clean = bg.convert('RGB')
        elif pil_img.mode != 'RGB':
            clean = pil_img.convert('RGB')
        else:
            clean = pil_img
        fd, path = tempfile.mkstemp(suffix='.jpg')
        os.close(fd)
        clean.save(path, format='JPEG', quality=95)
        run.add_picture(path, width=Inches(width_in))
        os.remove(path)

def add_cell_text(cell, text, size_pt, rgb, align):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*rgb)

photo_dir = r'C:\Users\Administrator\Desktop\新建文件夹'
photos = sorted([os.path.join(photo_dir, f) for f in os.listdir(photo_dir)
                 if f.lower().endswith(('.jpg','.jpeg','.png'))])
print(f'Photos: {len(photos)}')

# Portrait A4, 4 per page
for orient_name, is_landscape, ipp in [("纵_4张", False, 4), ("横_4张", True, 4)]:
    doc = Document()
    sec = doc.sections[0]

    # Set page size via XML
    pg_sz = sec._sectPr.find(qn('w:pgSz'))
    if is_landscape:
        pg_sz.set(qn('w:w'), str(int(29.7 / 2.54 * 1440)))
        pg_sz.set(qn('w:h'), str(int(21 / 2.54 * 1440)))
        pg_sz.set(qn('w:orient'), 'landscape')
    else:
        pg_sz.set(qn('w:w'), str(int(21 / 2.54 * 1440)))
        pg_sz.set(qn('w:h'), str(int(29.7 / 2.54 * 1440)))
        pg_sz.set(qn('w:orient'), 'portrait')

    pg_mar = sec._sectPr.find(qn('w:pgMar'))
    if pg_mar is None:
        pg_mar = OxmlElement('w:pgMar'); sec._sectPr.append(pg_mar)
    for attr, val in [('w:top', 1.1), ('w:bottom', 1.2), ('w:left', 2.4), ('w:right', 1.7)]:
        pg_mar.set(qn(attr), str(int(val / 2.54 * 1440)))

    # Cell width: available width / 2
    avail = 21 - 2.4 - 1.7
    max_w_in = avail / 2 / 2.54

    tp = doc.add_paragraph("估价对象现状照片")
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.runs[0]
    tr.font.size = Pt(12); tr.font.bold = True

    cap_align = WD_ALIGN_PARAGRAPH.CENTER
    rgb = (0x33, 0x33, 0x33)
    cap_pt = 10
    cap_tpl = "{f}"

    n = len(photos)
    pair_idx = 0
    while pair_idx * 2 < n:
        pairs_per_page = (ipp + 1) // 2
        page_end = min(pair_idx + pairs_per_page, (n + 1) // 2)
        for _ in range(pair_idx, page_end):
            left_i = pair_idx * 2; right_i = pair_idx * 2 + 1
            has_right = right_i < n

            img_row = doc.add_table(rows=1, cols=2)
            img_row.style = 'Table Grid'; img_row.alignment = 1
            set_cell_width(img_row.rows[0].cells[0], max_w_in)
            set_cell_width(img_row.rows[0].cells[1], max_w_in)
            add_cell_image(img_row.rows[0].cells[0], photos[left_i], max_w_in)
            if has_right:
                add_cell_image(img_row.rows[0].cells[1], photos[right_i], max_w_in)

            cap_row = doc.add_table(rows=1, cols=2)
            cap_row.style = 'Table Grid'; cap_row.alignment = 1
            set_cell_width(cap_row.rows[0].cells[0], max_w_in)
            set_cell_width(cap_row.rows[0].cells[1], max_w_in)
            lf = os.path.splitext(os.path.basename(photos[left_i]))[0]
            add_cell_text(cap_row.rows[0].cells[0], cap_tpl.replace("{n}", str(left_i+1)).replace("{f}", lf), cap_pt, rgb, cap_align)
            if has_right:
                rf = os.path.splitext(os.path.basename(photos[right_i]))[0]
                add_cell_text(cap_row.rows[0].cells[1], cap_tpl.replace("{n}", str(right_i+1)).replace("{f}", rf), cap_pt, rgb, cap_align)
            pair_idx += 1

    out = os.path.join(os.environ['TEMP'], f'gen_{orient_name}.docx')
    doc.save(out)
    print(f'{orient_name}: {os.path.getsize(out):,} bytes')

print("SUCCESS")
